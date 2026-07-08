from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"D:\2026\202605\hitrise")
ANDROID_RES = ROOT / "hitrise-android" / "app" / "src" / "main" / "res"
DOCS = ROOT / "hitrise-docs"
OUT = DOCS / "HitRise_用户手册.docx"

APP_ICON = ROOT / "hitrise-deploy" / "icon_preview" / "hitrise_app_icon_square_1024.png"
if not APP_ICON.exists():
    APP_ICON = ANDROID_RES / "mipmap-xxxhdpi" / "ic_launcher.png"
HOME_BANNER = ANDROID_RES / "drawable-nodpi" / "home_banner.png"
HOME_REPORT_BG = ANDROID_RES / "drawable-nodpi" / "home_report_bg.png"
HOME_ACHIEVEMENT_BG = ANDROID_RES / "drawable-nodpi" / "home_achievement_bg.png"


FONT = "Microsoft YaHei"
TEXT = "17343B"
MUTED = "557A7D"
TEAL = "10BDAA"
TEAL_DARK = "096D65"
TEAL_SOFT = "EFFFFA"
TEAL_LINE = "CDEFE8"
ORANGE = "FF8A32"
ORANGE_SOFT = "FFF1E4"
YELLOW = "FFD060"
RED = "E65A4F"
BLUE = "2E74B5"
LIGHT = "F7FFFD"
WHITE = "FFFFFF"


def set_font(run, size=None, bold=None, color=None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_run(paragraph, text, size=10.5, bold=False, color=TEXT):
    run = paragraph.add_run(text)
    set_font(run, size=size, bold=bold, color=color)
    return run


def set_para(paragraph, before=0, after=6, line=1.25, align=None):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if align is not None:
        paragraph.alignment = align


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def border(cell, color=TEAL_LINE, size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def margins(cell, top=90, start=120, bottom=90, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    width = tc_pr.first_child_found_in("w:tcW")
    if width is None:
        width = OxmlElement("w:tcW")
        tc_pr.append(width)
    width.set(qn("w:w"), str(dxa))
    width.set(qn("w:type"), "dxa")


def set_table_grid(table, widths):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    old_grid = tbl.tblGrid
    if old_grid is not None:
        tbl.remove(old_grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    tbl.insert(1, grid)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])


def style_table(table, widths, header=True):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_grid(table, widths)
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            shade(cell, TEAL_DARK if header and r_idx == 0 else WHITE)
            border(cell, "BFEFE5" if r_idx != 0 else "7FE5D8", "8")
            margins(cell, 100, 140, 100, 140)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx == 0 or (header and r_idx == 0) else WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                set_font(run, size=9.2, bold=(header and r_idx == 0), color=(WHITE if header and r_idx == 0 else TEXT))


def add_body(doc, text, after=6, color=TEXT, size=10.5):
    p = doc.add_paragraph()
    add_run(p, text, size=size, color=color)
    set_para(p, after=after)
    return p


def add_h1(doc, text):
    p = doc.add_paragraph()
    add_run(p, text, size=16, bold=True, color=TEAL_DARK)
    set_para(p, before=18, after=10, line=1.25)
    return p


def add_h2(doc, text):
    p = doc.add_paragraph()
    add_run(p, text, size=13, bold=True, color=BLUE)
    set_para(p, before=14, after=7, line=1.25)
    return p


def add_h3(doc, text):
    p = doc.add_paragraph()
    add_run(p, text, size=12, bold=True, color=TEAL_DARK)
    set_para(p, before=10, after=5, line=1.25)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        add_run(p, item, size=10.2, color=TEXT)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        add_run(p, item, size=10.2, color=TEXT)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)


def add_callout(doc, title, body, fill=TEAL_SOFT, line=TEAL):
    table = doc.add_table(rows=1, cols=1)
    style_table(table, [9360], header=False)
    cell = table.cell(0, 0)
    shade(cell, fill)
    border(cell, line, "10")
    margins(cell, 160, 180, 160, 180)
    p = cell.paragraphs[0]
    add_run(p, title, size=10.5, bold=True, color=TEAL_DARK)
    set_para(p, after=3)
    p2 = cell.add_paragraph()
    add_run(p2, body, size=9.8, color=TEXT)
    set_para(p2, after=0)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_table(doc, rows, widths):
    table = doc.add_table(rows=1, cols=len(rows[0]))
    for c, text in enumerate(rows[0]):
        p = table.rows[0].cells[c].paragraphs[0]
        add_run(p, text, size=9.2, bold=True, color=WHITE)
    for row_data in rows[1:]:
        row = table.add_row()
        for c, text in enumerate(row_data):
            p = row.cells[c].paragraphs[0]
            add_run(p, text, size=9.2, color=TEXT)
    style_table(table, widths, header=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(5)
    return table


def add_picture(doc, path, caption, width=5.9):
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(p, after=2)
    r = p.add_run()
    r.add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(cap, caption, size=8.8, color=MUTED)
    set_para(cap, after=8)


def setup_doc():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(10.2)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(footer, "HitRise 用户手册  |  本手册适用于当前 Hitrise Android 版本", size=8.5, color=MUTED)
    return doc


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(p, before=20, after=12)
    if APP_ICON.exists():
        p.add_run().add_picture(str(APP_ICON), width=Inches(1.45))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "HitRise", size=30, bold=True, color=TEAL_DARK)
    set_para(p, after=4)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "家庭健身 · 燃脂拳击", size=16, bold=True, color=TEXT)
    set_para(p, after=8)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "用户手册", size=24, bold=True, color=ORANGE)
    set_para(p, after=10)

    meta = doc.add_table(rows=5, cols=2)
    rows = [
        ("产品名称", "HitRise"),
        ("内部代号", "hitrise"),
        ("产品码", "HTR01"),
        ("蓝牙设备", "SENBALL# 系列智能拳击速度球"),
        ("适用版本", "当前 Hitrise Android 版本，applicationId: com.zclei.hitrise"),
    ]
    for r, (k, v) in enumerate(rows):
        add_run(meta.cell(r, 0).paragraphs[0], k, size=9.5, bold=True, color=TEAL_DARK)
        add_run(meta.cell(r, 1).paragraphs[0], v, size=9.5, color=TEXT)
    style_table(meta, [1900, 7460], header=False)

    add_callout(
        doc,
        "重要提示",
        "HitRise 训练数据、卡路里与等效燃脂量用于运动反馈和训练激励，不作为医疗、营养诊断或治疗建议。训练前请确认场地安全、设备牢固、电量充足。",
        fill=ORANGE_SOFT,
        line=ORANGE,
    )

    doc.add_section(WD_SECTION_START.NEW_PAGE)


def add_overview(doc):
    add_h1(doc, "1. 产品概览")
    add_body(
        doc,
        "HitRise 是面向家庭健身和燃脂拳击训练的智能运动 APP。APP 通过 SENBALL# 智能拳击速度球采集击打次数、击打力度、电量和训练状态，并在锻炼中心实时显示回合计时、拳数、BPM、卡路里、等效燃脂、连击识别和力度曲线。",
    )
    add_picture(doc, HOME_BANNER, "锻炼中心首页主视觉：家庭健身 · 燃脂拳击", width=5.8)
    add_table(
        doc,
        [
            ("模块", "主要用途", "用户能看到什么"),
            ("锻炼中心", "开始/结束训练、查看实时训练仪表盘", "回合计时、拳数、BPM、卡路里、等效燃脂、击打力度、连接状态/最新战报"),
            ("锻炼成果", "查看徽章与荣誉段位", "锻炼时间、拳击次数、最大力度、平均力度、卡路里、等效燃脂六类徽章"),
            ("榜单排名", "查看云端排名", "按锻炼时间、击拳数、最大力度、平均力度、卡路里、等效燃脂排名"),
            ("个人中心", "查看账号与累计数据", "个人资料、历史训练、云端同步状态、协议入口"),
        ],
        [1500, 3100, 4760],
    )


def add_quick_start(doc):
    add_h1(doc, "2. 快速开始")
    add_h2(doc, "2.1 首次使用")
    add_numbered(
        doc,
        [
            "安装并打开 HitRise，允许必要的蓝牙权限。Android 12 及以上系统需要允许附近设备权限。",
            "打开手机蓝牙，并保持 SENBALL# 设备有电、处于可连接状态。",
            "进入右上角齿轮设置，扫描并连接 SENBALL# 设备。APP 会优先使用免配对 BLE 连接，避免弹出系统蓝牙配对请求。",
            "连接成功后，返回锻炼中心，确认“连接状态/最新战报”区域显示设备名称、电量和已连接状态。",
            "点击“开始”，完成 3-2-1-GO 后进入实时训练。",
        ],
    )
    add_h2(doc, "2.2 每次训练前检查")
    add_bullets(
        doc,
        [
            "设备名称应以 SENBALL# 开头，且最后一位为英文字母。",
            "电量显示 0-100 表示实际电量；101 显示“充电”；102 显示“充满”。",
            "请确认速度球安装牢固，周围无易碎物品，身体状态适合训练。",
            "训练时建议让手机屏幕保持亮屏，便于观察实时数据和回合状态。",
        ],
    )


def add_home(doc):
    add_h1(doc, "3. 锻炼中心首页")
    add_body(
        doc,
        "锻炼中心首页采用浅绿色运动风格。顶部为主视觉图，右上角小齿轮进入主设置；训练卡片中间为回合计时环，左右为“开始”和“结束”按钮，打开首页即可看到核心训练入口。",
    )
    add_h2(doc, "3.1 实时训练卡片")
    add_table(
        doc,
        [
            ("区域", "说明"),
            ("回合标签", "显示当前第几回合以及总回合数，例如“第 1/3 回合”。"),
            ("计时环", "显示当前回合剩余时间或休息时间。训练中颜色会随状态变化。"),
            ("开始/结束", "开始用于进入倒计时和训练；结束用于提前结束当前训练并生成战报。"),
            ("四项指标", "总击打数、BPM、卡路里、等效燃脂。单位显示在数字下方并居中。"),
            ("今日目标", "默认 500 拳，实时显示今日已完成拳数与进度条。"),
        ],
        [2200, 7160],
    )
    add_h2(doc, "3.2 点击开始后的自动上移")
    add_body(
        doc,
        "点击“开始”后，APP 会自动将实时训练卡片移动到屏幕上端，便于完整查看实时训练与击打力度区域。第 2、3 回合开始时也会自动执行同样的上移动作。",
    )


def add_bluetooth(doc):
    add_h1(doc, "4. 蓝牙连接与电量")
    add_h2(doc, "4.1 设备筛选规则")
    add_body(
        doc,
        "HitRise 扫描时只显示符合规则的设备：名称以 SENBALL# 开头，且设备名最后一位是英文字母。例如 SENBALL#00000G 会被识别为有效设备。",
    )
    add_h2(doc, "4.2 连接步骤")
    add_numbered(
        doc,
        [
            "点击首页右上角齿轮，进入“蓝牙与语言设置”。",
            "点击“扫描”，等待出现 SENBALL# 设备。",
            "选中设备后点击“连接”。连接成功后，顶部状态和首页连接状态区域会更新。",
            "后续打开 APP 时会尝试自动连接上次使用过的 BLE 设备。",
        ],
    )
    add_callout(
        doc,
        "免配对说明",
        "当前版本保留打开 APP 时自动蓝牙连接，但会阻止 SENBALL# 设备进入系统配对流程。若系统仍弹出配对请求，请先取消请求，再回到 APP 重新扫描连接。",
    )
    add_h2(doc, "4.3 电量显示")
    add_table(
        doc,
        [
            ("设备返回值", "APP 显示", "含义"),
            ("0-100", "0%-100%", "实际电量百分比"),
            ("101", "充电", "设备正在充电"),
            ("102", "充满", "设备已充满"),
        ],
        [2000, 2200, 5160],
    )
    add_body(
        doc,
        "训练开始后，为避免蓝牙辅助刷新影响实时拳数，电量状态变化会暂缓刷新；训练结束后再恢复正常刷新。连接状态/最新战报和蓝牙设置等辅助区域约 5 秒刷新一次。",
        color=MUTED,
        size=9.8,
    )


def add_training_settings(doc):
    add_h1(doc, "5. 训练设置")
    add_body(
        doc,
        "实时训练卡片右上角“训练设置”用于配置回合、休息、训练方式和节拍速度。设置会保存到本机，下次进入训练时继续使用。",
    )
    add_table(
        doc,
        [
            ("预设", "训练时长", "休息时长", "回合数", "建议人群"),
            ("初学者", "1 分钟", "0.5 分钟", "3 回合", "默认选择，适合首次体验、家庭日常燃脂"),
            ("经典", "2 分钟", "0.5 分钟", "3 回合", "适合有一定运动基础的用户"),
            ("高强度", "5 分钟", "1 分钟", "3 回合", "适合需要更高训练量的用户"),
            ("HIIT", "1 分钟", "1 分钟", "6 回合", "适合间歇训练"),
        ],
        [1400, 1500, 1500, 1200, 3760],
    )
    add_h2(doc, "5.1 节拍与音效")
    add_bullets(
        doc,
        [
            "训练方式可选择自由参考或跟拍节奏，节拍速度支持 40-140 BPM。",
            "训练中使用拳击音效增强反馈；休息时播放轻松舒缓的休息音乐。",
            "当前主设置已取消用户自选云端音效、背景音乐和配色入口，系统保持当前默认方案。",
        ],
    )


def add_realtime_training(doc):
    add_h1(doc, "6. 实时训练")
    add_h2(doc, "6.1 倒计时与回合流程")
    add_numbered(
        doc,
        [
            "点击“开始”后，仪表盘中间显示 3-2-1-GO。",
            "GO 后如果蓝牙计数通道暂时未就绪，APP 会进入“计数通道准备中”，延迟重试或自动重连；通道确认成功后才正式进入训练计时。",
            "回合训练中，APP 根据蓝牙“数据2”增加拳击次数，并实时更新拳数、BPM、卡路里和等效燃脂。",
            "每个回合休息时，APP 会关闭陀螺仪计数；到达下一回合训练时间后重新打开陀螺仪。",
            "训练中若蓝牙偶发断开，APP 会自动尝试重连，减少复杂展会环境下的掉线影响。",
        ],
    )
    add_h2(doc, "6.2 击打力度")
    add_body(
        doc,
        "击打力度来自蓝牙“数据6”和“数据7”组合计算，单位为 N。APP 已按传感器识别数据统一乘以 0.6 后显示，以贴合当前设备标定。实时训练界面以折线图展示力度变化，并将本次最小力度视为坐标 0，使较大力度更容易看出波动。",
    )
    add_table(
        doc,
        [
            ("颜色", "含义"),
            ("浅绿色", "轻击或较低力度"),
            ("绿色", "中等力度"),
            ("黄色/橙色", "重拳或较高力度"),
            ("红色", "爆发力度"),
        ],
        [1800, 7560],
    )
    add_h2(doc, "6.3 连击识别与 AI 教练")
    add_bullets(
        doc,
        [
            "连击识别不区分左手/右手，也不区分直拳、勾拳、摆拳。",
            "APP 根据击打间隔识别重击、连击、三连击和爆发连击。",
            "AI 教练在回合开始、节奏变化、最后 10 秒等关键节点进行低频语音提示，避免频繁打断训练。",
            "AI 教练实时指导默认在后台运行，界面上不占用额外区域。",
        ],
    )


def add_calorie(doc):
    add_h1(doc, "7. 卡路里与等效燃脂")
    add_body(
        doc,
        "当前版本的卡路里计算不再只按拳数固定折算，而是综合训练时长、拳击频率和平均击打力度估算动态强度。等效燃脂量由卡路里换算而来，用于帮助用户理解训练消耗。",
    )
    add_table(
        doc,
        [
            ("参数", "规则"),
            ("默认体重", "70 kg"),
            ("基础拳击 MET", "约 7.8"),
            ("频率因子", "按每分钟拳数计算，并限制在 0.50-1.60"),
            ("力度因子", "按平均力度相对参考力度开方计算，并限制在 0.70-1.35"),
            ("动态 MET", "由频率 70% + 力度 30% 综合得到，并限制在合理运动范围"),
            ("卡路里", "动态 MET × 3.5 × 70 kg ÷ 200 × 训练分钟数"),
            ("等效燃脂", "卡路里 ÷ 7.7，单位 g"),
        ],
        [2200, 7160],
    )
    add_callout(
        doc,
        "估算口径",
        "卡路里和等效燃脂是训练反馈估算值。不同个体的体重、代谢、动作质量和运动基础不同，真实消耗会有差异。",
        fill=ORANGE_SOFT,
        line=ORANGE,
    )


def add_report(doc):
    add_h1(doc, "8. 训练战报与分享")
    add_body(
        doc,
        "每个回合结束后，APP 会生成回合训练战报。多回合训练会按回合累计锻炼时间、累计击拳数、卡路里和等效燃脂，并将每一回合结果纳入云端统计。",
    )
    add_picture(doc, HOME_REPORT_BG, "连接状态/最新战报背景素材", width=3.9)
    add_table(
        doc,
        [
            ("战报字段", "说明"),
            ("累计锻炼时间", "从第 1 回合开始累计到当前回合"),
            ("累计击拳数", "当前训练已完成回合的总拳数"),
            ("最大力度", "当前训练已记录的最高单拳力度"),
            ("平均力度", "当前训练已记录力度的平均值"),
            ("卡路里", "按当前累计训练强度估算"),
            ("等效燃脂", "按卡路里除以 7.7 估算"),
            ("平均 BPM/最佳连击", "用于评估训练节奏和爆发能力"),
        ],
        [2200, 7160],
    )
    add_h2(doc, "8.1 分享入口")
    add_bullets(
        doc,
        [
            "训练战报可生成分享图。",
            "锻炼成果可分享荣誉与徽章。",
            "榜单排名可分享当前排名。",
            "分享图已统一为当前首页色彩系统，避免旧版深色风格混杂。",
        ],
    )


def add_achievements(doc):
    add_h1(doc, "9. 锻炼成果")
    add_body(
        doc,
        "锻炼成果页面按六类指标展示徽章。徽章上方显示精简目标文字，徽章本身使用统一风格图片；解锁进度来自云端训练累计数据。",
    )
    add_picture(doc, HOME_ACHIEVEMENT_BG, "目标与成就背景素材", width=4.0)
    add_table(
        doc,
        [
            ("类别", "四档目标"),
            ("锻炼时间", "60 分钟、300 分钟、600 分钟、2000 分钟"),
            ("拳击次数", "100 拳、500 拳、1000 拳、5000 拳"),
            ("最大拳击力度", "500N、1000N、1300N、1600N"),
            ("平均拳击力度", "500N、800N、1000N、1200N"),
            ("卡路里消耗", "500 kcal、1000 kcal、2000 kcal、4000 kcal"),
            ("等效燃脂量", "100g、500g、1000g、2000g"),
        ],
        [2200, 7160],
    )
    add_h2(doc, "9.1 荣誉段位")
    add_body(
        doc,
        "荣誉段位用于展示用户阶段性成长，当前段位体系已按拳击训练重新整理，不再使用旧模式中“30 秒最佳”等残留口径。",
    )


def add_cloud(doc):
    add_h1(doc, "10. 云端同步、榜单与个人中心")
    add_h2(doc, "10.1 云端同步")
    add_body(
        doc,
        "完成训练后，APP 会把训练会话和每回合数据上传到云端。云端用于历史记录、锻炼成果、榜单排名和个人中心累计统计。若网络异常，APP 会保留本地训练结果，恢复网络后可再次同步。",
    )
    add_table(
        doc,
        [
            ("同步数据", "用途"),
            ("训练会话", "记录一次训练的累计时长、总拳数、卡路里、等效燃脂、平均 BPM、最大/平均力度"),
            ("回合数据", "记录每回合拳数、累计拳数、回合时长、卡路里、等效燃脂和力度指标"),
            ("成果统计", "用于徽章解锁、荣誉段位和个人中心累计数据"),
            ("榜单数据", "用于锻炼时间、拳数、力度、卡路里、等效燃脂排名"),
        ],
        [2200, 7160],
    )
    add_h2(doc, "10.2 榜单排名")
    add_bullets(
        doc,
        [
            "榜单支持按锻炼时间、累计击拳数、最大力度、平均力度、卡路里和等效燃脂量切换。",
            "若当前用户有排名，会在“我的排名”区域显示。",
            "榜单数据来自云端，刷新时需要网络连接。",
        ],
    )
    add_h2(doc, "10.3 个人中心")
    add_bullets(
        doc,
        [
            "查看个人资料、训练历史、累计数据和云端同步状态。",
            "支持头像设置。",
            "可查看隐私政策、用户协议和开发者联系信息。",
        ],
    )


def add_settings(doc):
    add_h1(doc, "11. 主设置、语言与协议")
    add_h2(doc, "11.1 主设置")
    add_body(
        doc,
        "首页右上角齿轮进入主设置。当前版本主设置保留蓝牙连接和语言选择；配色选择、云端音效选择、背景音乐选择入口已取消。系统保持当前 APP 配色、默认拳击音效和无训练背景音乐，但回合休息时仍播放舒缓休息音乐。",
    )
    add_h2(doc, "11.2 四国语言")
    add_table(
        doc,
        [
            ("语言", "说明"),
            ("中文", "默认主要说明语言"),
            ("English", "English UI text"),
            ("Français", "Interface française"),
            ("ไทย", "ภาษาไทย"),
        ],
        [2200, 7160],
    )
    add_h2(doc, "11.3 隐私政策与用户协议")
    add_body(
        doc,
        "APP 内置中文、英文、法文、泰文四种语言的隐私政策和用户协议。用户可在个人中心或关于页面查看。",
    )


def add_safety(doc):
    add_h1(doc, "12. 安全使用与维护")
    add_h2(doc, "12.1 训练安全")
    add_bullets(
        doc,
        [
            "训练前做好热身，训练后适当拉伸。",
            "请在平整、通风、无遮挡环境中训练。",
            "未成年人、老人或有心血管疾病、关节损伤等情况的用户，应在监护或专业建议下使用。",
            "如出现头晕、胸闷、关节疼痛等不适，请立即停止训练。",
        ],
    )
    add_h2(doc, "12.2 设备维护")
    add_bullets(
        doc,
        [
            "保持设备电量充足，充电时留意 APP 电量状态。",
            "避免设备进水、重摔或在强磁环境中长期放置。",
            "若蓝牙连接异常，可关闭手机蓝牙后重新开启，或在 APP 中断开后重新扫描连接。",
            "展会、商场等蓝牙信号复杂环境可能出现短时延迟，APP 会尽量自动重连并保持训练数据连续。",
        ],
    )


def add_faq(doc):
    add_h1(doc, "13. 常见问题")
    rows = [
        ("问题", "处理建议"),
        ("打开 APP 时出现蓝牙配对请求怎么办？", "点击取消。HitRise 使用免配对 BLE 连接，并已阻止 SENBALL# 进入配对流程。如仍反复出现，重启蓝牙后在 APP 设置中重新扫描。"),
        ("连接后没有立即显示电量怎么办？", "连接成功后 APP 会主动读取电量。训练中电量刷新会暂缓，训练结束后恢复正常刷新。"),
        ("训练中拳数突然跳变怎么办？", "复杂蓝牙环境下可能出现短时数据延迟。APP 已降低辅助区域刷新频率，并优先保证实时计数通道。"),
        ("GO 后没有进入训练怎么办？", "如果蓝牙写通道暂时未就绪，APP 会进入计数通道准备中并自动重试。请保持设备靠近手机。"),
        ("卡路里和等效燃脂为什么与其他设备不同？", "HitRise 按拳数、频率、力度和时长估算，不同设备算法不同，数据应用于训练反馈，不作为医学结论。"),
        ("休息时为什么有音乐？", "休息音乐用于区分训练和休息状态，帮助用户调整呼吸和节奏。"),
        ("分享图样式和页面不一致怎么办？", "当前版本已将训练战报、成果和榜单分享图统一到首页色彩系统。请确认安装的是最新 APK。"),
    ]
    add_table(doc, rows, [3000, 6360])


def add_appendix(doc):
    add_h1(doc, "14. 附录：关键数据口径")
    add_table(
        doc,
        [
            ("项目", "当前口径"),
            ("拳击次数", "蓝牙协议“数据2”增加逻辑"),
            ("拳击力度", "蓝牙“数据6/数据7”组合计算，显示值为传感器力度 × 0.6，单位 N"),
            ("电量", "0-100 显示百分比，101 显示充电，102 显示充满"),
            ("BPM", "按最近击打间隔估算训练节奏"),
            ("训练默认值", "初学者：1 分钟训练、0.5 分钟休息、3 回合"),
            ("云端入口", "APP API Base URL: https://hitrise.86086.cn/hitrise/api/v1/"),
        ],
        [2300, 7060],
    )
    add_body(doc, "文档版本：2026-06-28 重新生成。", color=MUTED, size=9.2)


def build():
    doc = setup_doc()
    add_cover(doc)
    add_overview(doc)
    add_quick_start(doc)
    add_home(doc)
    add_bluetooth(doc)
    add_training_settings(doc)
    add_realtime_training(doc)
    add_calorie(doc)
    add_report(doc)
    add_achievements(doc)
    add_cloud(doc)
    add_settings(doc)
    add_safety(doc)
    add_faq(doc)
    add_appendix(doc)
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build())
